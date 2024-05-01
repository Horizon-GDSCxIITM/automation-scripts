from docx import Document
from pypdf import PdfReader
import re

def extract_emails_and_numbers(pdf_path):
    # Open the PDF file in read-binary mode
    with open(pdf_path, 'rb') as file:
        pdf = PdfReader(file)

        # Initialize an empty string to store extracted text
        text = ""

        # Extract text from each page of the PDF
        for page_num in range(pdf.get_num_pages()):
            page = pdf.get_page(page_num)
            text += page.extract_text()

        # Use regex to find email addresses
        emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)

        # Use regex to find phone numbers (this is a basic pattern, adjust as needed)
        phone_numbers = re.findall(r'\b(?:\+\d{1,2}\s)?\d{10}\b', text)

        return emails, phone_numbers



def extract_emails_and_numbers_from_docx(uploaded_file):
    # Initialize an empty string to store extracted text
    text = ""

    # Extract text from the DOCX file
    doc = Document(uploaded_file)
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # Use regex to find email addresses
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)

    # Use regex to find phone numbers (this is a basic pattern, adjust as needed)
    phone_numbers = re.findall(r'\b(?:\+\d{1,2}\s)?\d{10}\b', text)

    return emails, phone_numbers
# # Example usage
# pdf_path = ''
# # emails, phone_numbers = extract_emails_and_numbers_from_docx(pdf_path)

# print("Emails found:")
# for email in emails:
#     print(email)

# print("\nPhone numbers found:")
# for phone_number in phone_numbers:
#     print(phone_number)

# from docx import Document
# import re

# def read_docx(filename):
#     doc = Document(filename)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# def apply_regex(text, pattern):
#     return re.findall(pattern, text)

# filename = 'heemSen.doc'
# doc_text = read_docx(filename)

# # Example regex pattern to find email addresses
# email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

# emails_found = apply_regex(doc_text, email_pattern)
# print(emails_found)

