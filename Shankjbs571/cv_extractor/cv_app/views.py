import os
from django.shortcuts import render
from docx import Document
from pypdf import PdfReader
import re
import pandas as pd
from django.http import HttpResponse
from .templates import *

# # Create your views here.

# def index(request):
#     return {"THIS IS FROM CV_APP"}


def index(request):
    return render(request,'base.html')

def upload_files(request):
    if request.method == 'POST' and request.FILES['file']:
        data = []
        uploaded_files = request.FILES.getlist('file')
        
        
        print(uploaded_files)
        for uploaded_file in uploaded_files:
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            if file_extension == '.pdf':
                emails, phone_numbers = extract_emails_and_numbers_from_pdf(uploaded_file)
            elif file_extension == '.docx':
                emails, phone_numbers = extract_emails_and_numbers_from_docx(uploaded_file)
                print("this is docx")
            
            if emails:
                print("Emails found:")
                for email in emails:
                    print(email)
            if phone_numbers:
                print("\nPhone numbers found:")
            #     for phone_number in phone_numbers:
            #         print(phone_number)
            # with open(uploaded_file.name, 'wb+') as destination:
            #     for chunk in uploaded_file.chunks():
            #         destination.write(chunk)

            # Append file name, emails, and phone numbers to the data list
            data.append({
                'File Name': uploaded_file.name,
                'Emails': '\n'.join(emails),
                'Phone Numbers': '\n'.join(phone_numbers)
            })

        # Create a DataFrame from the data list
        df = pd.DataFrame(data)
        # Write the DataFrame to an Excel file
        output_path = 'output.xlsx'
        df.to_excel(output_path, index=False)

        # Serve the Excel file as a download
        with open(output_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = 'attachment; filename=output.xlsx'
            return response
        # return HttpResponse('Files uploaded successfully')
    return render(request, 'your_template.html')




def extract_emails_and_numbers_from_pdf(file):
    pdf = PdfReader(file)

    # Initialize an empty string to store extracted text
    text = ""

    # Extract text from each page of the PDF
    for page_num in range(pdf.get_num_pages()):
        page = pdf.get_page(page_num)
        text += page.extract_text()

    # Use regex to find email addresses
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)

    # Use regex to find phone numbers (this is a basic pattern, adjust as needed)
    phone_numbers = re.findall(r'\b(?:\+\d{1,2}\s?)?\d\s?\d\s?\d\s?\d\s?\d\s?\d\s?\d\s?\d\s?\d\s?\d\b', text)


    return emails, phone_numbers


def extract_emails_and_numbers_from_docx(uploaded_file):
    # Initialize an empty string to store extracted text
    text = ""

    # Extract text from the DOCX file
    doc = Document(uploaded_file)
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # Use regex to find email addresses
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)

    # Use regex to find phone numbers (this is a basic pattern, adjust as needed)
    phone_numbers = re.findall(r'\b(?:\+\d{1,2}\s)?\d{10}\b', text)

    return emails, phone_numbers



# from docx import Document
# import re

# def read_docx(filename):
#     doc = Document(filename)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)

# def apply_regex(text, pattern):
#     return re.findall(pattern, text)

# filename = 'heemSen.doc'
# doc_text = read_docx(filename)

# Example regex pattern to find email addresses
# email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

# emails_found = apply_regex(doc_text, email_pattern)
# print(emails_found)